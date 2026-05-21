"""
Nexus QA — Doküman Parse Modülü
PDF, DOCX, TXT, MD formatlarını metin + metadata'ya çevirir.
Büyük dokümanlar chunk'lara bölünür (3000 token ~= 12000 karakter).

Desteklenen formatlar:
  .pdf   — pdfplumber (tablo + metin)
  .docx  — python-docx
  .txt   — düz metin
  .md    — Markdown (başlık yapısı korunur)
"""
from __future__ import annotations

import io
import logging
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import BinaryIO

logger = logging.getLogger(__name__)

# Chunk boyutu: ~3000 token ≈ 12000 karakter (GPT tokenizer ortalama 4 char/token)
CHUNK_SIZE_CHARS = 12_000
CHUNK_OVERLAP_CHARS = 500   # Bağlam sürekliliği için örtüşme


@dataclass
class ParsedDocument:
    """Parse edilmiş doküman."""
    filename: str
    format: str                          # pdf | docx | txt | md
    full_text: str = ""                  # Tüm metin
    chunks: list[str] = field(default_factory=list)  # Chunk'lara bölünmüş
    page_count: int = 0
    word_count: int = 0
    char_count: int = 0
    sections: list[str] = field(default_factory=list)  # Başlıklar (varsa)
    tables: list[list[list[str]]] = field(default_factory=list)  # PDF tablolar
    error: str | None = None

    def is_empty(self) -> bool:
        return not self.full_text.strip()

    def needs_chunking(self) -> bool:
        return self.char_count > CHUNK_SIZE_CHARS

    def summary(self) -> str:
        return (
            f"Dosya: {self.filename} | Format: {self.format} | "
            f"Sayfa: {self.page_count} | Kelime: {self.word_count} | "
            f"Chunk: {len(self.chunks)}"
        )


# ── Parse Fonksiyonları ──────────────────────────────────────────────────────

def _parse_pdf(data: bytes, filename: str) -> ParsedDocument:
    """PDF parse — pdfplumber ile."""
    try:
        import pdfplumber
    except ImportError:
        raise ImportError(
            "pdfplumber yüklü değil. "
            "pip install pdfplumber --break-system-packages"
        )

    doc = ParsedDocument(filename=filename, format="pdf")
    texts = []
    tables_all = []

    with pdfplumber.open(io.BytesIO(data)) as pdf:
        doc.page_count = len(pdf.pages)
        for page in pdf.pages:
            # Tablo çıkar
            for tbl in page.extract_tables() or []:
                if tbl:
                    tables_all.append(tbl)
                    # Tabloyu metne de ekle (basit format)
                    for row in tbl:
                        row_text = " | ".join(str(c or "").strip() for c in row)
                        if row_text.strip():
                            texts.append(row_text)

            # Normal metin
            text = page.extract_text()
            if text:
                texts.append(text.strip())

    doc.tables = tables_all
    doc.full_text = "\n\n".join(texts)
    return doc


def _parse_docx(data: bytes, filename: str) -> ParsedDocument:
    """DOCX parse — python-docx ile."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx yüklü değil. "
            "pip install python-docx --break-system-packages"
        )

    doc_obj = Document(io.BytesIO(data))
    doc = ParsedDocument(filename=filename, format="docx")
    texts = []
    sections = []

    for para in doc_obj.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        # Başlık tespiti
        if para.style.name.startswith("Heading"):
            sections.append(text)
            texts.append(f"\n## {text}\n")
        else:
            texts.append(text)

    # Tablolar
    tables_all = []
    for table in doc_obj.tables:
        tbl = []
        table_texts = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            tbl.append(cells)
            table_texts.append(" | ".join(cells))
        tables_all.append(tbl)
        texts.append("\n".join(table_texts))

    doc.sections = sections
    doc.tables = tables_all
    doc.full_text = "\n\n".join(texts)
    doc.page_count = 1  # DOCX sayfa sayısı hesaplanamaz kolay
    return doc


def _parse_txt(data: bytes, filename: str, fmt: str = "txt") -> ParsedDocument:
    """TXT / MD parse."""
    doc = ParsedDocument(filename=filename, format=fmt)
    # Encoding tespiti — önce UTF-8, sonra latin-1
    for enc in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
        try:
            doc.full_text = data.decode(enc)
            break
        except UnicodeDecodeError:
            continue
    else:
        doc.full_text = data.decode("utf-8", errors="replace")

    # MD ve TXT için başlıkları çıkar
    if fmt in ("md", "txt"):
        # Markdown-style headers (# Heading) — both md and txt files can have these
        sections = re.findall(r"^#{1,3}\s+(.+)$", doc.full_text, re.MULTILINE)
        if not sections:
            # Fallback: all-caps lines or lines ending with ":" as section headers
            sections = re.findall(r"^([A-ZÇŞĞÜÖİ][A-ZÇŞĞÜÖİ\s]{3,}|.+:)\s*$", doc.full_text, re.MULTILINE)
            sections = [s.rstrip(":").strip() for s in sections if len(s.strip()) > 3][:20]
        doc.sections = sections

    doc.page_count = max(1, len(doc.full_text) // 3000)
    return doc


# ── Chunk Bölme ─────────────────────────────────────────────────────────────

def _split_into_chunks(text: str) -> list[str]:
    """
    Metni CHUNK_SIZE_CHARS boyutunda parçalara böl.
    Paragraf sınırlarına göre böler (kelime kesmez).
    """
    if len(text) <= CHUNK_SIZE_CHARS:
        return [text]

    chunks: list[str] = []
    paragraphs = text.split("\n\n")
    current_chunk: list[str] = []
    current_len = 0

    for para in paragraphs:
        para_len = len(para)

        if current_len + para_len > CHUNK_SIZE_CHARS and current_chunk:
            chunks.append("\n\n".join(current_chunk))
            # Overlap: son paragrafı yeni chunk'a ekle
            overlap = current_chunk[-1:] if current_chunk else []
            current_chunk = overlap
            current_len = sum(len(p) for p in current_chunk)

        current_chunk.append(para)
        current_len += para_len

    if current_chunk:
        chunks.append("\n\n".join(current_chunk))

    logger.info(f"Doküman {len(chunks)} chunk'a bölündü (toplam {len(text)} karakter)")
    return chunks


# ── Ana Parse Fonksiyonu ─────────────────────────────────────────────────────

def parse_document(
    data: bytes,
    filename: str,
    content_type: str | None = None,
) -> ParsedDocument:
    """
    Dosya verisini ve adını alır, ParsedDocument döndürür.
    Format; dosya uzantısı ve/veya content_type'dan otomatik belirlenir.
    """
    ext = Path(filename).suffix.lower().lstrip(".")
    if not ext and content_type:
        ct = content_type.lower()
        if "pdf" in ct:
            ext = "pdf"
        elif "docx" in ct or "word" in ct or "openxml" in ct:
            ext = "docx"
        elif "markdown" in ct:
            ext = "md"
        else:
            ext = "txt"

    try:
        if ext == "pdf":
            doc = _parse_pdf(data, filename)
        elif ext == "docx":
            doc = _parse_docx(data, filename)
        elif ext == "md":
            doc = _parse_txt(data, filename, fmt="md")
        else:  # txt + diğerleri
            doc = _parse_txt(data, filename, fmt="txt")
    except Exception as exc:
        logger.error(f"Doküman parse hatası ({filename}): {exc}")
        doc = ParsedDocument(filename=filename, format=ext or "unknown")
        doc.error = str(exc)
        return doc

    # İstatistikler
    doc.word_count = len(doc.full_text.split())
    doc.char_count = len(doc.full_text)

    # Chunk'lara böl
    doc.chunks = _split_into_chunks(doc.full_text)

    logger.info(f"Parse tamamlandı: {doc.summary()}")
    return doc


# ── Token sayacı (yaklaşık) ──────────────────────────────────────────────────

def estimate_tokens(text: str) -> int:
    """Yaklaşık token sayısı (GPT: ortalama 4 karakter/token)."""
    return len(text) // 4


def truncate_for_ai(text: str, max_tokens: int = 3000) -> str:
    """AI'ya gönderilecek metni token limitine göre kes."""
    max_chars = max_tokens * 4
    if len(text) <= max_chars:
        return text
    # Son tam cümleye kadar kes
    truncated = text[:max_chars]
    last_period = max(
        truncated.rfind(". "),
        truncated.rfind(".\n"),
        truncated.rfind("\n\n"),
    )
    if last_period > max_chars // 2:
        truncated = truncated[:last_period + 1]
    return truncated + "\n\n[... doküman kesildi, devamı sonraki chunk'ta ...]"
