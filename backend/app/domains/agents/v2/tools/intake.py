"""Universal Intake — PDF/DOCX/URL/Swagger/Postman → normalize metin."""
from __future__ import annotations

import json
import logging
import os
import re
from pathlib import Path
from typing import Any

import httpx

logger = logging.getLogger(__name__)


class IntakeError(RuntimeError):
    pass


def parse_pdf(path: str | Path) -> str:
    path = Path(path)
    if not path.exists():
        raise IntakeError(f"PDF bulunamadı: {path}")
    try:
        import pdfplumber
        chunks: list[str] = []
        with pdfplumber.open(path) as pdf:
            for i, page in enumerate(pdf.pages):
                txt = page.extract_text() or ""
                if txt.strip():
                    chunks.append(f"--- Sayfa {i + 1} ---\n{txt}")
        return "\n\n".join(chunks)
    except ImportError:
        pass
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        return "\n\n".join(
            f"--- Sayfa {i + 1} ---\n{page.extract_text() or ''}"
            for i, page in enumerate(reader.pages)
        )
    except ImportError as exc:
        raise IntakeError("PDF parse için pdfplumber veya pypdf gerekli") from exc


def parse_docx(path: str | Path) -> str:
    path = Path(path)
    if not path.exists():
        raise IntakeError(f"DOCX bulunamadı: {path}")
    try:
        from docx import Document
    except ImportError as exc:
        raise IntakeError("DOCX parse için python-docx gerekli") from exc

    doc = Document(str(path))
    chunks: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            prefix = ""
            if p.style and p.style.name.startswith("Heading"):
                try:
                    level = int(p.style.name.replace("Heading ", ""))
                except (ValueError, AttributeError):
                    level = 1
                prefix = "#" * level + " "
            chunks.append(prefix + p.text)
    for i, table in enumerate(doc.tables):
        chunks.append(f"\n--- Tablo {i + 1} ---")
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            chunks.append(" | ".join(cells))
    return "\n".join(chunks)


async def parse_url(url: str, timeout: float = 10.0) -> str:
    try:
        async with httpx.AsyncClient(follow_redirects=True, timeout=timeout) as client:
            resp = await client.get(url, headers={"User-Agent": "TestwrightAI-Intake/1.0"})
            resp.raise_for_status()
            html = resp.text
    except httpx.HTTPError as exc:
        raise IntakeError(f"URL alınamadı: {exc}") from exc

    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "aside", "svg"]):
            tag.decompose()
        title = soup.title.string.strip() if soup.title and soup.title.string else ""
        body_text = soup.get_text(separator="\n", strip=True)
        return f"# {title}\n\nURL: {url}\n\n{body_text}"
    except ImportError:
        plain = re.sub(r"<[^>]+>", " ", html)
        plain = re.sub(r"\s+", " ", plain).strip()
        return f"URL: {url}\n\n{plain[:20_000]}"


def parse_openapi(spec_path_or_json: str) -> str:
    p = Path(spec_path_or_json)
    if p.exists():
        raw = p.read_text(encoding="utf-8")
    else:
        raw = spec_path_or_json
    try:
        spec = json.loads(raw)
    except json.JSONDecodeError:
        try:
            import yaml
            spec = yaml.safe_load(raw)
        except Exception as exc:
            raise IntakeError(f"OpenAPI parse hatası: {exc}") from exc

    lines: list[str] = []
    info = spec.get("info", {})
    lines.append(f"# {info.get('title', 'API')} v{info.get('version', '?')}")
    if info.get("description"):
        lines.append(info["description"])
    lines.append("")

    for s in spec.get("servers", []):
        lines.append(f"Server: {s.get('url', '')}")
    lines.append("")

    paths = spec.get("paths", {})
    lines.append(f"## Endpointler ({len(paths)} adet)")
    for path, methods in paths.items():
        for method, op in methods.items():
            if method not in {"get", "post", "put", "patch", "delete"}:
                continue
            lines.append(f"\n### {method.upper()} {path}")
            if op.get("summary"):
                lines.append(f"  {op['summary']}")
            if op.get("operationId"):
                lines.append(f"  operationId: {op['operationId']}")
            if op.get("tags"):
                lines.append(f"  tags: {', '.join(op['tags'])}")
            for prm in op.get("parameters", []):
                req = " (required)" if prm.get("required") else ""
                lines.append(f"    - {prm.get('name')} ({prm.get('in')}){req}")
            if op.get("requestBody"):
                lines.append("  Body: var")
            if op.get("responses"):
                lines.append(f"  Yanıtlar: {', '.join(op['responses'].keys())}")

    return "\n".join(lines)


def parse_postman(collection_path_or_json: str) -> str:
    p = Path(collection_path_or_json)
    if p.exists():
        raw = p.read_text(encoding="utf-8")
    else:
        raw = collection_path_or_json
    try:
        coll = json.loads(raw)
    except json.JSONDecodeError as exc:
        raise IntakeError(f"Postman parse: {exc}") from exc

    lines: list[str] = []
    info = coll.get("info", {})
    lines.append(f"# Postman: {info.get('name', '?')}")
    if info.get("description"):
        lines.append(str(info["description"])[:500])
    lines.append("")

    def walk(items: list[dict], depth: int = 0) -> None:
        for item in items:
            name = item.get("name", "")
            if "item" in item:
                lines.append(f"{'  ' * depth}## {name}")
                walk(item["item"], depth + 1)
            elif "request" in item:
                req = item["request"]
                method = req.get("method", "GET")
                url = req.get("url", {})
                url_str = url if isinstance(url, str) else url.get("raw", "")
                lines.append(f"{'  ' * depth}- {method} {url_str} — {name}")

    walk(coll.get("item", []))
    return "\n".join(lines)


async def intake_to_text(
    source: str,
    source_type: str | None = None,
    *,
    extra_context: str | None = None,
) -> tuple[str, str]:
    stype = (source_type or "").lower()
    if not stype:
        stype = _infer_source_type(source)

    if stype == "pdf":
        text = parse_pdf(source)
    elif stype == "docx":
        text = parse_docx(source)
    elif stype == "url":
        text = await parse_url(source)
    elif stype in {"swagger", "openapi"}:
        text = parse_openapi(source)
        stype = "swagger"
    elif stype == "postman":
        text = parse_postman(source)
    elif stype in {"text", "manual"}:
        text = source
    else:
        p = Path(source) if os.sep in source or source.endswith((".pdf", ".docx", ".json", ".yaml", ".yml", ".txt", ".md")) else None
        if p and p.exists():
            text = p.read_text(encoding="utf-8", errors="ignore")
            stype = "text"
        else:
            text = source
            stype = "text"

    if extra_context:
        text = text + "\n\n---\nEK BAĞLAM:\n" + extra_context

    return text, stype


def _infer_source_type(source: str) -> str:
    s = source.strip()
    if s.startswith(("http://", "https://")):
        return "url"
    p = Path(s) if os.sep in s or s.endswith((".pdf", ".docx", ".json", ".yaml", ".yml", ".txt", ".md")) else None
    if p and p.exists():
        suf = p.suffix.lower()
        if suf == ".pdf":
            return "pdf"
        if suf == ".docx":
            return "docx"
        if suf == ".json":
            try:
                data = json.loads(p.read_text(encoding="utf-8", errors="ignore"))
                if "info" in data and "item" in data:
                    return "postman"
                if "openapi" in data or "swagger" in data:
                    return "swagger"
            except Exception:
                pass
            return "swagger"
        if suf in (".yaml", ".yml"):
            return "swagger"
        if suf in (".txt", ".md"):
            return "text"
    return "text"
